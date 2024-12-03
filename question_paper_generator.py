import os
from docx import Document
import random

# Parse Question Bank Document
def parse_question_bank(file_path):
    document = Document(file_path)
    questions = []
    for para in document.paragraphs:
        if para.text.strip():  # Ignore blank lines
            questions.append(para.text.strip())
    return questions

# Categorize Questions
def categorize_questions(questions):
    categorized = {"unit_1": [], "unit_2": [], "unit_3": [], "unit_4": [], "unit_5": []}
    for q in questions:
        # Basic classification based on keywords (for demo purposes)
        if "unit 1" in q.lower():
            categorized["unit_1"].append(q)
        elif "unit 2" in q.lower():
            categorized["unit_2"].append(q)
        elif "unit 3" in q.lower():
            categorized["unit_3"].append(q)
        elif "unit 4" in q.lower():
            categorized["unit_4"].append(q)
        elif "unit 5" in q.lower():
            categorized["unit_5"].append(q)
    return categorized

# Select Questions
def select_questions(categorized):
    part_a = []
    part_b = []

    for unit, questions in categorized.items():
        if len(questions) >= 2:
            part_a.extend(random.sample(questions, 2))  # Select 2 questions for Part A
        if len(questions) >= 1:
            part_b.extend(random.sample(questions, 1))  # Select 1 question for Part B (with choice)

    return part_a, part_b

# Generate Question Paper Document
def generate_question_paper(part_a, part_b, output_file):
    document = Document()
    document.add_heading("Question Paper", level=1)

    document.add_heading("Part A - 10 Marks", level=2)
    for i, question in enumerate(part_a, 1):
        document.add_paragraph(f"{i}. {question}")

    document.add_heading("Part B - 50 Marks", level=2)
    for i, question in enumerate(part_b, 1):
        document.add_paragraph(f"{i}. {question} (OR)")

    document.save(output_file)
    print(f"Question Paper saved as {output_file}")

# Main
if __name__ == "__main__":
    input_file = "question_bank.docx"  
    output_file = "final_question_paper.docx"

    print("Parsing question bank...")
    questions = parse_question_bank(input_file)

    print("Categorizing questions...")
    categorized = categorize_questions(questions)

    print("Selecting questions...")
    part_a, part_b = select_questions(categorized)

    print("Generating question paper...")
    generate_question_paper(part_a, part_b, output_file)
    print("Done!")
